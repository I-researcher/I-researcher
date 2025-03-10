import os
import re
import asyncio
import pandas as pd
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass
from io import BytesIO
from docx import Document

import streamlit as st
from dotenv import load_dotenv

from langchain_core.messages import HumanMessage
from langchain_groq import ChatGroq

# Load environment variables
load_dotenv()


# -------------------------------------------------------------------
# 1. Model Configuration
# -------------------------------------------------------------------
@dataclass
class ModelConfig:
    name: str
    api_key: str
    model_id: str
    max_tokens: int


#GROQ_API_KEY = os.getenv("GROQ_API_KEY", "YOUR_GROQ_API_KEY")
GROQ_API_KEY='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'

DEEPSEEK_CONFIG = ModelConfig(
    name="DeepSeek",
    api_key=GROQ_API_KEY,
    model_id="deepseek-r1-distill-llama-70b",
    max_tokens=8192
)


# -------------------------------------------------------------------
# 2. Paper Manager Class
# -------------------------------------------------------------------
class PaperManager:
    def __init__(self):
        self.initialize_session_state()
        self.executor = ThreadPoolExecutor(max_workers=5)
        self.max_retries = 3

    def initialize_session_state(self):
        """Initialize Streamlit session state variables."""
        if 'settings' not in st.session_state:
            st.session_state.settings = {
                'temperature': 0.7,
                'word_counts': {},
                'min_word_count': 300
            }
        if 'scientific_paper' not in st.session_state:
            st.session_state.scientific_paper = {
                'title': "",
                'references': "",
                'sections': {section: "" for section in [
                    'Abstract', 'Introduction', 'Literature Review',
                    'Methodology', 'Results', 'Discussion', 'Conclusion'
                ]}
            }

    def create_model_instance(self):
        """Create a DeepSeek model instance with API key."""
        return ChatGroq(
            groq_api_key=DEEPSEEK_CONFIG.api_key,
            model_name=DEEPSEEK_CONFIG.model_id,
            temperature=st.session_state.settings['temperature'],
            max_tokens=DEEPSEEK_CONFIG.max_tokens
        )

    async def generate_section(self, prompt: str, min_words: int) -> str:
        """Generate high-quality scientific content with retries."""
        model = self.create_model_instance()
        for attempt in range(self.max_retries):
            try:
                loop = asyncio.get_event_loop()
                response = await loop.run_in_executor(
                    self.executor,
                    lambda: model.invoke([HumanMessage(content=prompt)])
                )
                result = response.content.strip()
                word_count = len(result.split())

                # Remove AI "thinking" process from DeepSeek
                result = clean_deepseek_output(result)

                # Ensure output meets the word limit
                if word_count < min_words:
                    return f"Error: The output is too short ({word_count} words). Increase token limit or refine the prompt."

                # Track word count
                st.session_state.settings['word_counts']['DeepSeek'] = word_count

                return result
            except Exception as e:
                if attempt == self.max_retries - 1:
                    return f"Error after {self.max_retries} attempts: {str(e)}"
                await asyncio.sleep(2)
        return "Unexpected error occurred"


def clean_deepseek_output(text: str) -> str:
    """Removes everything between <Think> and </Think>, including the tags."""
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()


# -------------------------------------------------------------------
# 3. File Upload & Data Processing
# -------------------------------------------------------------------
def load_uploaded_file(uploaded_file):
    """Load and process a CSV or Excel file dynamically."""
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith('.xlsx'):
                df = pd.read_excel(uploaded_file)
            else:
                st.error("Unsupported file format! Upload a CSV or Excel file.")
                return None

            # Show preview in sidebar
            st.sidebar.write("### File Preview:")
            st.sidebar.dataframe(df.head(5))

            # Show available columns
            st.sidebar.write("### Detected Features:")
            st.sidebar.write(df.columns.tolist())

            return df
        except Exception as e:
            st.error(f"Error loading file: {str(e)}")
    return None


# -------------------------------------------------------------------
# 4. Word File Handling
# -------------------------------------------------------------------
def save_to_word():
    """Saves the generated content to a Word file."""
    doc = Document()

    doc.add_heading(st.session_state.scientific_paper['title'], level=1)

    for section, content in st.session_state.scientific_paper['sections'].items():
        if content.strip():
            doc.add_heading(section, level=2)
            doc.add_paragraph(content)

    # Save to an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# -------------------------------------------------------------------
# 5. Streamlit Application
# -------------------------------------------------------------------
def main():
    st.set_page_config(page_title="Scientific Paper Assistant", layout="wide", page_icon="📝")

    paper_manager = PaperManager()

    st.title("Scientific Paper Assistant 📝")

    st.session_state.scientific_paper['title'] = st.text_input("Paper Title",
                                                               value=st.session_state.scientific_paper['title'])

    selected_section = st.selectbox("Select Section", list(st.session_state.scientific_paper['sections'].keys()))

    # File Upload Section
    uploaded_file = st.sidebar.file_uploader("Upload Results (CSV/XLSX)", type=["csv", "xlsx"])
    data_summary = ""

    if uploaded_file:
        data = load_uploaded_file(uploaded_file)
        if data is not None:
            data_summary = data.describe().to_string()

    min_words = st.slider("Select Minimum Word Count", min_value=100, max_value=1000, step=50,
                          value=st.session_state.settings['min_word_count'])

    user_instructions = st.text_area("Instructions")

    if st.button("Generate / Refine Section"):
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        prompt = f"""
        Title: {st.session_state.scientific_paper['title']}
        Section: {selected_section}
        Instructions: {user_instructions}
        Data Summary from Uploaded File:
        {data_summary}
        Please generate at least {min_words} words.
        """

        final_text = loop.run_until_complete(paper_manager.generate_section(prompt, min_words))
        st.session_state.scientific_paper['sections'][selected_section] = final_text
        st.success(f"Updated {selected_section}! Length: {len(final_text.split())} words")

    # **Dynamic Width Adjustment for Generated Content**
    text_length = len(st.session_state.scientific_paper['sections'][selected_section])
    text_area_height = min(600, max(200, text_length // 5))  # Adjusts height based on text length

    st.subheader(f"Generated {selected_section} Section")
    st.text_area("Generated Content",
                 value=st.session_state.scientific_paper['sections'][selected_section],
                 height=text_area_height)

    # **Save and Download Button**
    if st.button("Save to Word & Download"):
        buffer = save_to_word()
        st.download_button(
            label="Download Word Document",
            data=buffer,
            file_name="Scientific_Paper.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


if __name__ == "__main__":
    main()
